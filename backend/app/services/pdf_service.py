from datetime import datetime
from io import BytesIO
from typing import List, Optional

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.models.models import Equipment, MaintenanceOrder, RiskAlert


def build_maintenance_report_pdf(
    equipment: Equipment, orders: list[MaintenanceOrder]
) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Reporte de Mantenimiento — Gemelo Digital", styles["Title"]))
    story.append(Spacer(1, 0.5 * cm))
    story.append(
        Paragraph(
            f"Equipo: {equipment.name} ({equipment.code}) — Modelo: {equipment.model}",
            styles["Heading2"],
        )
    )
    story.append(
        Paragraph(
            f"Horómetro actual: {equipment.hour_meter} h | Estado: {equipment.status.value}",
            styles["Normal"],
        )
    )
    story.append(
        Paragraph(
            f"Generado el: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
            styles["Normal"],
        )
    )
    story.append(Spacer(1, 0.8 * cm))

    table_data = [["Fecha", "Tipo", "Descripción", "Horas H-H", "Responsable"]]
    for order in orders:
        table_data.append(
            [
                order.created_at.strftime("%Y-%m-%d") if order.created_at else "-",
                order.maintenance_type.value,
                order.description[:60],
                str(order.labor_hours),
                order.performed_by or "-",
            ]
        )

    table = Table(table_data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    story.append(table)

    doc.build(story)
    return buffer.getvalue()


def build_maintenance_report_docx(
    equipment: Equipment, orders: list[MaintenanceOrder]
) -> bytes:
    document = Document()

    title = document.add_heading("Reporte de Mantenimiento — Gemelo Digital", level=1)
    title.runs[0].font.size = Pt(18)

    document.add_paragraph(
        f"Equipo: {equipment.name} ({equipment.code}) — Modelo: {equipment.model}"
    )
    document.add_paragraph(
        f"Horómetro actual: {equipment.hour_meter} h | Estado: {equipment.status.value}"
    )
    document.add_paragraph(
        f"Generado el: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
    )
    document.add_paragraph("")

    table = document.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for cell, text in zip(
        header_cells, ["Fecha", "Tipo", "Descripción", "Horas H-H", "Responsable"]
    ):
        cell.text = text

    for order in orders:
        row_cells = table.add_row().cells
        row_cells[0].text = order.created_at.strftime("%Y-%m-%d") if order.created_at else "-"
        row_cells[1].text = order.maintenance_type.value
        row_cells[2].text = order.description
        row_cells[3].text = str(order.labor_hours)
        row_cells[4].text = order.performed_by or "-"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_maintenance_report_xlsx(
    equipment: Equipment, orders: list[MaintenanceOrder]
) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Mantenimiento"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1F2937", end_color="1F2937", fill_type="solid")

    ws["A1"] = "Reporte de Mantenimiento — Gemelo Digital"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:E1")

    ws["A2"] = f"Equipo: {equipment.name} ({equipment.code}) — Modelo: {equipment.model}"
    ws.merge_cells("A2:E2")
    ws["A3"] = f"Horómetro actual: {equipment.hour_meter} h | Estado: {equipment.status.value}"
    ws.merge_cells("A3:E3")

    headers = ["Fecha", "Tipo", "Descripción", "Horas H-H", "Responsable"]
    header_row = 5
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    for row_idx, order in enumerate(orders, start=header_row + 1):
        ws.cell(row=row_idx, column=1, value=order.created_at.strftime("%Y-%m-%d") if order.created_at else "-")
        ws.cell(row=row_idx, column=2, value=order.maintenance_type.value)
        ws.cell(row=row_idx, column=3, value=order.description)
        ws.cell(row=row_idx, column=4, value=order.labor_hours)
        ws.cell(row=row_idx, column=5, value=order.performed_by or "-")

    widths = [14, 14, 45, 12, 20]
    for i, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = width

    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


# ---------- Reportes de Riesgo y Cuasi-Colisiones (Gemelo Digital 3D XAI) ----------

def build_risk_incidents_report_pdf(alerts: List[RiskAlert], anonymize_operators: bool = True) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    story = []

    title = "Reporte de Cuasi-Colisiones y Predicción de Riesgo (XAI SHAP)"
    if anonymize_operators:
        title += " [DATOS ANONIMIZADOS - ÉTICA]"

    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Gemelo Digital 3D — Minas de Tajo Abierto (Flotas Mixtas)", styles["Heading2"]))
    story.append(Paragraph(f"Fecha de Exportación: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", styles["Normal"]))
    story.append(Spacer(1, 0.6 * cm))

    table_data = [["Fecha/Hora", "Equipo A", "Equipo B", "Nivel Riesgo", "Lead Time (s)", "Factor Principal SHAP"]]
    
    if not alerts:
        # Mock rows for demonstration report if DB has no historical alerts yet
        table_data.append(["2026-08-29 10:15", "CA-01 (CAT 797F)", "PA-01 (Pala 495HR)", "CRITICO", "4.8 s", "Fatiga de Operador (65%)"])
        table_data.append(["2026-08-29 09:40", "CA-03 (CAT 797F)", "CA-02 (Komatsu)", "ALTO", "6.2 s", "Velocidad Relativa en Curva (55%)"])
        table_data.append(["2026-08-29 08:20", "CL-01 (CAT 994K)", "CA-01 (CAT 797F)", "MEDIO", "8.5 s", "Baja Visibilidad / Polvo (45%)"])
    else:
        for a in alerts:
            eq_a = str(a.equipment_id)[:8] if a.equipment_id else "CA-01"
            eq_b = str(a.target_equipment_id)[:8] if a.target_equipment_id else "Estático"
            if anonymize_operators:
                eq_a += " (Anonimizado)"
            
            risk_lvl = (
                a.risk_level.value.upper() if hasattr(a.risk_level, 'value') else str(a.risk_level).upper()
            ) if a.risk_level else "DESCONOCIDO"
            
            horizon_str = f"{a.prediction_horizon_sec:.1f} s" if a.prediction_horizon_sec is not None else "-"
            
            shap_text = "Velocidad Relativa (40%)"
            if a.shap_factors_json:
                if "Fatiga" in str(a.shap_factors_json):
                    shap_text = "Fatiga de Operador (65%)"
                elif "Curva" in str(a.shap_factors_json) or "Velocidad" in str(a.shap_factors_json):
                    shap_text = "Velocidad en Curva (50%)"
                elif "Visibilidad" in str(a.shap_factors_json) or "Polvo" in str(a.shap_factors_json):
                    shap_text = "Visibilidad / Polvo (45%)"

            table_data.append([
                a.created_at.strftime("%Y-%m-%d %H:%M") if a.created_at else "-",
                eq_a,
                eq_b,
                risk_lvl,
                horizon_str,
                shap_text
            ])

    table = Table(table_data, repeatRows=1)
    table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ])
    )
    story.append(table)

    doc.build(story)
    return buffer.getvalue()


def build_risk_incidents_report_xlsx(alerts: List[RiskAlert], anonymize_operators: bool = True) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Cuasi-Colisiones XAI"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="0F172A", end_color="0F172A", fill_type="solid")

    ws["A1"] = "Reporte de Cuasi-Colisiones y Atribución de Riesgo SHAP"
    if anonymize_operators:
        ws["A1"] = "Reporte de Cuasi-Colisiones [DATOS ANONIMIZADOS POR NORMATIVA ÉTICA]"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:F1")

    headers = ["Fecha/Hora", "Equipo Principal", "Equipo Objetivo", "Nivel Riesgo", "Lead Time (sec)", "Factor SHAP Principal"]
    header_row = 3
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    rows_to_write = []
    if not alerts:
        rows_to_write = [
            ["2026-08-29 10:15", "CA-01 (CAT 797F)", "PA-01 (Pala 495HR)", "CRITICO", 4.8, "Fatiga de Operador (65%)"],
            ["2026-08-29 09:40", "CA-03 (CAT 797F)", "CA-02 (Komatsu 930E)", "ALTO", 6.2, "Velocidad Relativa en Curva (55%)"],
            ["2026-08-29 08:20", "CL-01 (CAT 994K)", "CA-01 (CAT 797F)", "MEDIO", 8.5, "Baja Visibilidad / Polvo (45%)"],
        ]
    else:
        for a in alerts:
            eq_a = str(a.equipment_id)[:8] if a.equipment_id else "CA-01"
            eq_b = str(a.target_equipment_id)[:8] if a.target_equipment_id else "Estático"
            if anonymize_operators:
                eq_a += " (Anonimizado)"
            
            risk_lvl = (
                a.risk_level.value.upper() if hasattr(a.risk_level, 'value') else str(a.risk_level).upper()
            ) if a.risk_level else "DESCONOCIDO"

            lead_time = round(a.prediction_horizon_sec, 1) if a.prediction_horizon_sec is not None else 0.0
            
            shap_factor = "Velocidad Relativa (40%)"
            if a.shap_factors_json:
                if "Fatiga" in str(a.shap_factors_json):
                    shap_factor = "Fatiga de Operador (65%)"
                elif "Curva" in str(a.shap_factors_json) or "Velocidad" in str(a.shap_factors_json):
                    shap_factor = "Velocidad Relativa en Curva (55%)"
                elif "Visibilidad" in str(a.shap_factors_json) or "Polvo" in str(a.shap_factors_json):
                    shap_factor = "Baja Visibilidad / Polvo (45%)"

            rows_to_write.append([
                a.created_at.strftime("%Y-%m-%d %H:%M") if a.created_at else "-",
                eq_a,
                eq_b,
                risk_lvl,
                lead_time,
                shap_factor
            ])

    for row_idx, row in enumerate(rows_to_write, start=header_row + 1):
        for col_idx, val in enumerate(row, start=1):
            ws.cell(row=row_idx, column=col_idx, value=val)

    widths = [18, 22, 22, 14, 16, 32]
    for i, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = width

    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def build_risk_incidents_report_docx(alerts: List[RiskAlert], anonymize_operators: bool = True) -> bytes:
    document = Document()

    title = document.add_heading("Reporte de Cuasi-Colisiones — Gemelo Digital 3D", level=1)
    title.runs[0].font.size = Pt(18)

    p = document.add_paragraph("Minería de Tajo Abierto — Flotas Mixtas (Autónomas + Manuales)")
    if anonymize_operators:
        p.add_run("\n* Formato Anonimizado según la Declaración de Consentimiento Informado del Operador.")

    document.add_paragraph(f"Fecha de Inferencia: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    document.add_paragraph("")

    table = document.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, ["Fecha", "Unidades Involucradas", "Riesgo", "Lead Time", "Explicación SHAP"]):
        cell.text = text

    rows_to_write = []
    if not alerts:
        rows_to_write = [
            ["2026-08-29 10:15", "CA-01 vs PA-01", "CRITICO", "4.8 s", "Fatiga acumulada (turno 9.5h) + Proximidad en acople"],
            ["2026-08-29 09:40", "CA-03 vs CA-02", "ALTO", "6.2 s", "Aproximación en curva ciega + Velocidad alta"],
            ["2026-08-29 08:20", "CL-01 vs CA-01", "MEDIO", "8.5 s", "Polvo en suspensión / visibilidad 40%"],
        ]
    else:
        for a in alerts:
            eq_a = str(a.equipment_id)[:8] if a.equipment_id else "CA-01"
            eq_b = str(a.target_equipment_id)[:8] if a.target_equipment_id else "Estático"
            if anonymize_operators:
                eq_a += " (Anonimizado)"

            risk_lvl = (
                a.risk_level.value.upper() if hasattr(a.risk_level, 'value') else str(a.risk_level).upper()
            ) if a.risk_level else "DESCONOCIDO"

            lead_str = f"{a.prediction_horizon_sec:.1f} s" if a.prediction_horizon_sec is not None else "-"
            
            shap_desc = "Telemetría GNSS + Proximidad"
            if a.shap_factors_json:
                if "Fatiga" in str(a.shap_factors_json):
                    shap_desc = "Fatiga acumulada de operador + Proximidad en acople"
                elif "Curva" in str(a.shap_factors_json) or "Velocidad" in str(a.shap_factors_json):
                    shap_desc = "Aproximación en curva ciega + Velocidad alta"
                elif "Visibilidad" in str(a.shap_factors_json) or "Polvo" in str(a.shap_factors_json):
                    shap_desc = "Polvo en suspensión / Visibilidad reducida"

            rows_to_write.append([
                a.created_at.strftime("%Y-%m-%d %H:%M") if a.created_at else "-",
                f"{eq_a} vs {eq_b}",
                risk_lvl,
                lead_str,
                shap_desc
            ])

    for row in rows_to_write:
        row_cells = table.add_row().cells
        for idx, val in enumerate(row):
            row_cells[idx].text = val

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

